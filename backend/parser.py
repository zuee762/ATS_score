"""Multi-format Resume Parser.

Supports PDF, DOCX, and TXT resumes.
Extracts contact info, summary, experience, skills, projects, education, and certifications.
"""

from typing import Dict, List, Any, Optional
import io
import re
from pypdf import PdfReader
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts raw text from PDF bytes."""
    reader = PdfReader(io.BytesIO(file_bytes))
    full_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text.append(text)
    return "\n".join(full_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts raw text from DOCX bytes."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables text if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))
    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extracts text from TXT bytes trying common encodings."""
    for enc in ["utf-8", "latin-1", "utf-16", "cp1252"]:
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def parse_resume_bytes(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Parses resume binary content into structured data according to file extension.
    """
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raw_text = extract_text_from_txt(file_bytes)

    return parse_resume_text(raw_text, filename)


def parse_resume_text(raw_text: str, filename: str = "resume.txt") -> Dict[str, Any]:
    """
    Extracts structured fields, contact information, and segmented sections from raw resume text.
    """
    # Normalize newlines
    clean_text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [l.strip() for l in clean_text.split("\n") if l.strip()]
    
    # 1. Contact Information Extraction
    email = extract_email(clean_text)
    phone = extract_phone(clean_text)
    linkedin = extract_linkedin(clean_text)
    github = extract_github(clean_text)
    candidate_name = extract_name(lines, email)
    
    # 2. Section Segmentation
    sections = segment_sections(clean_text)
    
    # 3. Specific Sub-structure extraction
    skills = parse_skills_section(sections.get("skills", ""))
    experience_items = parse_experience_section(sections.get("experience", ""))
    education_items = parse_education_section(sections.get("education", ""))
    projects_items = parse_projects_section(sections.get("projects", ""))
    certifications_items = parse_certifications_section(sections.get("certifications", ""))
    summary = sections.get("summary", "").strip()

    word_count = len(re.findall(r'\b\w+\b', clean_text))
    
    return {
        "filename": filename,
        "raw_text": clean_text,
        "word_count": word_count,
        "contact": {
            "name": candidate_name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "github": github
        },
        "sections": {
            "summary": summary,
            "skills": skills,
            "experience": experience_items,
            "education": education_items,
            "projects": projects_items,
            "certifications": certifications_items,
            "raw_sections": sections
        }
    }


def extract_email(text: str) -> Optional[str]:
    """Extracts first email address found."""
    match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extracts phone number in various formats."""
    match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    return match.group(0) if match else None


def extract_linkedin(text: str) -> Optional[str]:
    """Extracts LinkedIn profile URL or handle."""
    match = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+', text, re.IGNORECASE)
    return match.group(0) if match else None


def extract_github(text: str) -> Optional[str]:
    """Extracts GitHub profile URL."""
    match = re.search(r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+', text, re.IGNORECASE)
    return match.group(0) if match else None


def extract_name(lines: List[str], email: Optional[str]) -> str:
    """Heuristic to extract candidate name from top lines."""
    for line in lines[:5]:
        # Avoid lines that contain emails, phone numbers, or typical section titles
        if email and email in line:
            continue
        if any(w in line.lower() for w in ["resume", "curriculum", "page", "phone", "email", "linkedin", "http"]):
            continue
        # If line contains 2-4 words with letters and spaces only
        words = line.split()
        if 1 <= len(words) <= 4 and all(re.match(r'^[A-Za-z.\'-]+$', w) for w in words):
            return line.strip()
    return lines[0] if lines else "Candidate Name"


SECTION_PATTERNS = {
    "summary": r'(?:professional\s+summary|summary|profile|about\s+me|career\s+objective|objective)',
    "experience": r'(?:work\s+experience|professional\s+experience|experience|employment\s+history|work\s+history)',
    "skills": r'(?:technical\s+skills|skills\s+&?\s*competencies|skills\s+and\s+tools|core\s+competencies|skills|technologies)',
    "education": r'(?:education\s+&?\s*qualifications|academic\s+background|education|academics)',
    "projects": r'(?:key\s+projects|personal\s+projects|academic\s+projects|projects)',
    "certifications": r'(?:certifications\s+&?\s*licenses|certifications|certificates|licenses|achievements)'
}


def segment_sections(text: str) -> Dict[str, str]:
    """
    Segments resume text into canonical section keys.
    """
    lines = text.split("\n")
    sections: Dict[str, List[str]] = {}
    current_section = "header"
    sections[current_section] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check if line matches a major section header
        detected_section = None
        # Must be relatively short and prominent
        if len(stripped) < 40:
            clean_header = re.sub(r'[^a-zA-Z\s&]', '', stripped).strip()
            for key, pattern in SECTION_PATTERNS.items():
                if re.fullmatch(pattern, clean_header, re.IGNORECASE):
                    detected_section = key
                    break

        if detected_section:
            current_section = detected_section
            if current_section not in sections:
                sections[current_section] = []
        else:
            sections[current_section].append(stripped)

    # Join lines per section
    result = {}
    for k, v in sections.items():
        result[k] = "\n".join(v)
    return result


def parse_skills_section(skills_text: str) -> List[str]:
    """Parses individual skill keywords from raw skills section text."""
    if not skills_text:
        return []
    
    # Split by commas, bullets, colons, newlines, pipes, slashes
    delimiters = r'[,•|\n;:\t\/]+'
    raw_tokens = re.split(delimiters, skills_text)
    
    clean_skills = []
    stopwords = {"and", "with", "using", "proficient", "familiar", "knowledge", "expert", "intermediate", "advanced", "skills", "tools", "technologies", "languages"}
    
    for token in raw_tokens:
        tok = token.strip()
        tok = re.sub(r'^[-*•\d.]+', '', tok).strip()
        if 2 <= len(tok) <= 35 and tok.lower() not in stopwords:
            # Avoid long sentences
            if len(tok.split()) <= 4:
                clean_skills.append(tok)
                
    # Deduplicate while preserving order
    seen = set()
    deduped = []
    for s in clean_skills:
        if s.lower() not in seen:
            seen.add(s.lower())
            deduped.append(s)
    return deduped


def parse_experience_section(exp_text: str) -> List[Dict[str, Any]]:
    """Parses work experience blocks into structured items."""
    if not exp_text:
        return []

    lines = [l.strip() for l in exp_text.split("\n") if l.strip()]
    items = []
    current_item = None
    
    # Date regex pattern (e.g., "Jan 2022 - Present", "2020 - 2023", "05/2019 – 08/2021")
    date_pattern = r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|\d{2}/\d{4})\s*(?:-|–|to)\s*(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|Present|Current|\d{2}/\d{4})'

    for line in lines:
        has_date = re.search(date_pattern, line, re.IGNORECASE)
        # Bullet detection
        is_bullet = bool(re.match(r'^[-*•–—►▪]\s*', line)) or (current_item is not None and not has_date and len(line) > 50)
        
        if has_date or (not is_bullet and (current_item is None or len(current_item["bullets"]) >= 2)):
            # New role entry
            if current_item and (current_item["role"] or current_item["company"] or current_item["bullets"]):
                items.append(current_item)
                
            date_match = re.search(date_pattern, line, re.IGNORECASE)
            dates = date_match.group(0) if date_match else ""
            line_without_date = re.sub(date_pattern, '', line, flags=re.IGNORECASE).strip(' |,-')
            
            # Split role and company if formatted like "Software Engineer | Google"
            parts = [p.strip() for p in re.split(r'\|| - | at |, ', line_without_date) if p.strip()]
            role = parts[0] if len(parts) > 0 else "Role"
            company = parts[1] if len(parts) > 1 else ""
            
            current_item = {
                "role": role,
                "company": company,
                "dates": dates,
                "bullets": []
            }
        else:
            clean_bullet = re.sub(r'^[-*•–—►▪]\s*', '', line).strip()
            if current_item is not None and clean_bullet:
                current_item["bullets"].append(clean_bullet)

    if current_item and (current_item["role"] or current_item["company"] or current_item["bullets"]):
        items.append(current_item)
        
    return items


def parse_education_section(edu_text: str) -> List[Dict[str, Any]]:
    """Parses education items."""
    if not edu_text:
        return []
    lines = [l.strip() for l in edu_text.split("\n") if l.strip()]
    items = []
    for line in lines:
        clean = re.sub(r'^[-*•]\s*', '', line).strip()
        if len(clean) > 5:
            items.append({"details": clean})
    return items


def parse_projects_section(proj_text: str) -> List[Dict[str, Any]]:
    """Parses project items."""
    if not proj_text:
        return []
    lines = [l.strip() for l in proj_text.split("\n") if l.strip()]
    projects = []
    current_proj = None

    for line in lines:
        is_bullet = bool(re.match(r'^[-*•–—►▪]\s*', line)) or (current_proj and len(line) > 60)
        if not is_bullet and len(line) < 60:
            if current_proj:
                projects.append(current_proj)
            current_proj = {"title": line.strip(" -:*•"), "description": []}
        else:
            clean = re.sub(r'^[-*•–—►▪]\s*', '', line).strip()
            if current_proj:
                current_proj["description"].append(clean)
            else:
                current_proj = {"title": "Key Project", "description": [clean]}
                
    if current_proj:
        projects.append(current_proj)
    return projects


def parse_certifications_section(cert_text: str) -> List[str]:
    """Parses certifications list."""
    if not cert_text:
        return []
    lines = [l.strip() for l in cert_text.split("\n") if l.strip()]
    certs = []
    for line in lines:
        clean = re.sub(r'^[-*•]\s*', '', line).strip()
        if len(clean) > 3:
            certs.append(clean)
    return certs
