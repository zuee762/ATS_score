"""Resume Exporter.

Generates ATS-optimized PDF (via ReportLab) and Word DOCX (via python-docx) documents
from structured tailored resume data.
"""

from typing import Dict, List, Any
import io
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT


def generate_pdf_bytes(tailored_data: Dict[str, Any]) -> bytes:
    """
    Renders an ATS-friendly, single-column clean PDF resume using ReportLab.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    
    # Custom ATS-friendly styles
    name_style = ParagraphStyle(
        'ResumeName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#0f172a')
    )
    
    contact_style = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#475569')
    )
    
    section_heading_style = ParagraphStyle(
        'ResumeHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#1e293b'),
        spaceBefore=8,
        spaceAfter=3
    )
    
    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor('#0f172a')
    )
    
    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#334155')
    )
    
    bullet_style = ParagraphStyle(
        'ResumeBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#1e293b'),
        leftIndent=12,
        spaceAfter=2
    )

    story = []
    
    # 1. Contact Information Header
    contact = tailored_data.get("contact", {})
    name = contact.get("name") or "Professional Candidate"
    story.append(Paragraph(name.upper(), name_style))
    story.append(Spacer(1, 4))
    
    contact_items = []
    if contact.get("email"):
        contact_items.append(contact["email"])
    if contact.get("phone"):
        contact_items.append(contact["phone"])
    if contact.get("linkedin"):
        contact_items.append(contact["linkedin"])
    if contact.get("github"):
        contact_items.append(contact["github"])
        
    if contact_items:
        story.append(Paragraph(" • ".join(contact_items), contact_style))
        story.append(Spacer(1, 8))

    # 2. Professional Summary
    summary = tailored_data.get("summary")
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=6))
        story.append(Paragraph(summary, body_style))
        story.append(Spacer(1, 8))

    # 3. Technical Skills
    skills = tailored_data.get("skills", [])
    if skills:
        story.append(Paragraph("TECHNICAL & CORE SKILLS", section_heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=6))
        skills_text = " • ".join(skills)
        story.append(Paragraph(f"<b>Key Proficiencies:</b> {skills_text}", body_style))
        story.append(Spacer(1, 8))

    # 4. Work Experience
    experience = tailored_data.get("experience", [])
    if experience:
        story.append(Paragraph("WORK EXPERIENCE", section_heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=6))
        
        for exp in experience:
            role = exp.get("role", "Software Professional")
            company = exp.get("company", "Company")
            dates = exp.get("dates", "")
            
            header_text = f"<b>{role}</b>"
            if company:
                header_text += f" | {company}"
            if dates:
                header_text += f" <font color='#64748b'>({dates})</font>"
                
            story.append(Paragraph(header_text, job_title_style))
            story.append(Spacer(1, 3))
            
            for bullet in exp.get("bullets", []):
                story.append(Paragraph(f"• {bullet}", bullet_style))
                
            story.append(Spacer(1, 6))

    # 5. Projects
    projects = tailored_data.get("projects", [])
    if projects:
        story.append(Paragraph("PROJECTS", section_heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=6))
        for proj in projects:
            title = proj.get("title", "Key Project")
            story.append(Paragraph(f"<b>{title}</b>", job_title_style))
            for desc in proj.get("description", []):
                story.append(Paragraph(f"• {desc}", bullet_style))
            story.append(Spacer(1, 4))

    # 6. Education & Certifications
    education = tailored_data.get("education", [])
    certs = tailored_data.get("certifications", [])
    if education or certs:
        story.append(Paragraph("EDUCATION & CERTIFICATIONS", section_heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=6))
        
        for edu in education:
            details = edu.get("details", "") if isinstance(edu, dict) else str(edu)
            story.append(Paragraph(f"• {details}", bullet_style))
            
        for cert in certs:
            cert_name = cert if isinstance(cert, str) else cert.get("name", "")
            story.append(Paragraph(f"• <b>Certified:</b> {cert_name}", bullet_style))
            
        story.append(Spacer(1, 6))

    doc.build(story)
    return buffer.getvalue()


def generate_docx_bytes(tailored_data: Dict[str, Any]) -> bytes:
    """
    Renders an ATS-friendly editable DOCX document using python-docx.
    """
    doc = docx.Document()
    
    # Page Margins (0.5 inch all sides)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    contact = tailored_data.get("contact", {})
    name = contact.get("name") or "Professional Candidate"
    
    # Name Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = name_p.add_run(name.upper())
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(18)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(15, 23, 42)
    name_p.paragraph_format.space_after = Pt(2)

    # Contact line
    contact_items = []
    if contact.get("email"):
        contact_items.append(contact["email"])
    if contact.get("phone"):
        contact_items.append(contact["phone"])
    if contact.get("linkedin"):
        contact_items.append(contact["linkedin"])
    if contact.get("github"):
        contact_items.append(contact["github"])

    if contact_items:
        cont_p = doc.add_paragraph()
        cont_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cont = cont_p.add_run(" | ".join(contact_items))
        run_cont.font.name = 'Calibri'
        run_cont.font.size = Pt(9.5)
        run_cont.font.color.rgb = RGBColor(71, 85, 105)
        cont_p.paragraph_format.space_after = Pt(10)

    # Helper function for Section Headings
    def add_section_header(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

    # Summary
    summary = tailored_data.get("summary")
    if summary:
        add_section_header("PROFESSIONAL SUMMARY")
        p = doc.add_paragraph()
        r = p.add_run(summary)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # Skills
    skills = tailored_data.get("skills", [])
    if skills:
        add_section_header("TECHNICAL & CORE SKILLS")
        p = doc.add_paragraph()
        r_bold = p.add_run("Key Proficiencies: ")
        r_bold.font.name = 'Calibri'
        r_bold.font.bold = True
        r_bold.font.size = Pt(10)
        r = p.add_run(" • ".join(skills))
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # Experience
    experience = tailored_data.get("experience", [])
    if experience:
        add_section_header("WORK EXPERIENCE")
        for exp in experience:
            role = exp.get("role", "Software Professional")
            company = exp.get("company", "Company")
            dates = exp.get("dates", "")
            
            p_title = doc.add_paragraph()
            r_role = p_title.add_run(role)
            r_role.font.name = 'Calibri'
            r_role.font.bold = True
            r_role.font.size = Pt(10.5)
            
            if company:
                r_comp = p_title.add_run(f" | {company}")
                r_comp.font.name = 'Calibri'
                r_comp.font.size = Pt(10)
                
            if dates:
                r_dates = p_title.add_run(f" ({dates})")
                r_dates.font.name = 'Calibri'
                r_dates.font.italic = True
                r_dates.font.size = Pt(9.5)
                r_dates.font.color.rgb = RGBColor(100, 116, 139)
                
            p_title.paragraph_format.space_before = Pt(4)
            p_title.paragraph_format.space_after = Pt(2)
            
            for bullet in exp.get("bullets", []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                r_b = p_bullet.add_run(bullet)
                r_b.font.name = 'Calibri'
                r_b.font.size = Pt(9.5)
                p_bullet.paragraph_format.space_after = Pt(1)

    # Projects
    projects = tailored_data.get("projects", [])
    if projects:
        add_section_header("KEY PROJECTS")
        for proj in projects:
            title = proj.get("title", "Project")
            p_proj = doc.add_paragraph()
            r_p = p_proj.add_run(title)
            r_p.font.name = 'Calibri'
            r_p.font.bold = True
            r_p.font.size = Pt(10)
            
            for desc in proj.get("description", []):
                p_b = doc.add_paragraph(style='List Bullet')
                r_b = p_b.add_run(desc)
                r_b.font.name = 'Calibri'
                r_b.font.size = Pt(9.5)
                p_b.paragraph_format.space_after = Pt(1)

    # Education & Certs
    education = tailored_data.get("education", [])
    certs = tailored_data.get("certifications", [])
    if education or certs:
        add_section_header("EDUCATION & CERTIFICATIONS")
        for edu in education:
            details = edu.get("details", "") if isinstance(edu, dict) else str(edu)
            p_e = doc.add_paragraph(style='List Bullet')
            r_e = p_e.add_run(details)
            r_e.font.name = 'Calibri'
            r_e.font.size = Pt(9.5)
            
        for cert in certs:
            cname = cert if isinstance(cert, str) else cert.get("name", "")
            p_c = doc.add_paragraph(style='List Bullet')
            r_bold = p_c.add_run("Certified: ")
            r_bold.font.bold = True
            r_c = p_c.add_run(cname)
            r_c.font.name = 'Calibri'
            r_c.font.size = Pt(9.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
